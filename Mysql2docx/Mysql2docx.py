#!/usr/bin/python
# -*-coding:utf-8-*-

import json
import pymysql
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from entity import Table
from entity import Column

__author__ = 'skydu'


class Mysql2docx(object):
    db_name = ''

    #
    def __init__(self):
        db_name = ''

    @staticmethod
    def get_comment(comment):
        if comment is None:
            return ""
        try:
            data = json.loads(comment)
            return data[0]['value']
        except:
            return comment

    def get_tables(self, db):
        sql = "SELECT TABLE_NAME, TABLE_COMMENT FROM INFORMATION_SCHEMA.TABLES " \
              "WHERE TABLE_SCHEMA = '%s' AND TABLE_TYPE = 'BASE TABLE'" % self.db_name
        cursor = db.cursor()
        cursor.execute(sql)
        data = cursor.fetchall()
        tables = list()
        for table in data:
            t = Table(table[0], self.get_comment(table[1]))
            tables.append(t)
        cursor.close()
        return tables

    def get_columns(self, db, table_name):
        sql = "SELECT  " \
              "COLUMN_NAME 列名,  " \
              "COLUMN_TYPE 数据类型,  " \
              "IS_NULLABLE 是否为空,    " \
              "COLUMN_DEFAULT 默认值,    " \
              "COLUMN_COMMENT 备注   " \
              "FROM INFORMATION_SCHEMA.COLUMNS  " \
              "WHERE TABLE_SCHEMA ='%s' AND TABLE_NAME  = '%s';" % (self.db_name, table_name)
        cursor = db.cursor()
        cursor.execute(sql)
        data = cursor.fetchall()
        columns = list()
        for column in data:
            c = Column(column[0], column[1], column[2], column[3], self.get_comment(column[4]))
            columns.append(c)
        cursor.close()
        return columns

    def do(self, db_host, db_name, db_port=3306, db_user='root', db_password='root', file_name='数据字典'):
        print("jdbc_url=%s:%d/%s" % (db_host, db_port, db_name))
        print("db_user=%s" % db_user)
        print("db_pwd=%s" % db_password)
        print()

        self.db_name = db_name
        db = pymysql.connect(db_host, db_user, db_password, db_name, db_port, charset="utf8")
        tables = self.get_tables(db)
        for table in tables:
            table_name = table.name
            table.columns = self.get_columns(db, table_name)

        # make a document
        document = Document()
        document.styles['Normal'].font.name = u'宋体'  # 可换成word里面任意字体

        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(file_name)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.name = u'宋体'

        table_no = 0
        for table in tables:
            # print("table: %s columns: %d" % (table, len(table.columns)))

            document.add_page_break()
            table_no += 1

            document.add_heading(text="%d.%s %s" % (table_no, table.name, table.comment), level=1)
            # document.add_paragraph(text="%s %s" % (table.name, table.comment), style='List Number')

            trs = len(table.columns) + 1
            tds = 6
            t = document.add_table(rows=trs, cols=tds, style='Table Grid')

            # 设置行高
            for i in range(trs):
                tr = t.rows[i]._tr
                trPr = tr.get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), "400")
                trPr.append(trHeight)

            # 设置列宽
            t.columns[0].width = Inches(0.5)
            t.columns[1].width = Inches(1.5)
            t.columns[2].width = Inches(1.1)
            t.columns[3].width = Inches(2)
            t.columns[4].width = Inches(0.9)
            t.columns[5].width = Inches(1)

            # 设置表格头
            th_text = [u'序号', u'字段名', u'类型', u'字段描述', u'是否为空', u'默认值']
            for i in range(len(th_text)):
                p = t.rows[0].cells[i].paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(th_text[i])
                run.font.bold = True
                run.font.size = Pt(12)

            # 设置表格体
            i = 0
            for column in table.columns:
                i += 1
                '''
                for j in range(cols):
                    p = t.rows[i].cells[j].paragraph[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p.add_run(table_headers[i])
                    run.font.bold = True
                    run.font.size = Pt(12)
                '''
                '''
                rowCells = t.rows[i].cells
                rowCells[0].text = str(i)
                rowCells[1].text = column.name
                rowCells[2].text = column.type
                rowCells[3].text = column.comment
                rowCells[4].text = column.allowNull
                if column.defaultValue is not None:
                    rowCells[5].text = column.defaultValue
                '''
                rowCells = t.rows[i].cells
                # 序号
                p = rowCells[0].paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.add_run(str(i))

                rowCells[1].text = column.name
                rowCells[2].text = column.type
                rowCells[3].text = column.comment

                # 是否为空
                p = rowCells[4].paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.add_run(column.allowNull)

                if column.defaultValue is not None:
                    rowCells[5].text = column.defaultValue

        #
        document.save('%s.docx' % file_name)
