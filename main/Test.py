<<<<<<< HEAD
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('D:/Projects/res/material/911.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')

# m = Mysql2Docx()
# m.do(db_host='127.0.0.1', db_name='test', db_port=3307, db_user='root', db_pwd='root', file_name='数据字典_mysql', file_path='/')

# m = SqlServer2Docx()
# m.do(db_host='127.0.0.1', db_name='test', db_port=1433, db_user='root', db_pwd='root', file_name='数据字典_sqlServer', file_path='/')

# o = Oracle2Docx()
# o.do(db_host='127.0.0.1', db_name='test', db_port=3307, db_user='root', db_pwd='root', file_name='数据字典_oracle', file_path='/')
=======
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('D:/Projects/res/material/911.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')

# m = Mysql2Docx()
# m.do(db_host='127.0.0.1', db_name='test', db_port=3307, db_user='root', db_pwd='root', file_name='数据字典_mysql', file_path='/')

# m = SqlServer2Docx()
# m.do(db_host='127.0.0.1', db_name='test', db_port=1433, db_user='root', db_pwd='root', file_name='数据字典_sqlServer', file_path='/')

# o = Oracle2Docx()
# o.do(db_host='127.0.0.1', db_name='test', db_port=3307, db_user='root', db_pwd='root', file_name='数据字典_oracle', file_path='/')
>>>>>>> 1bf17071848bae1a6e563518c8bf3ba4807cb073
