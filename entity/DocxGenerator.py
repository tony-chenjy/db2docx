#!/usr/bin/python
# -*-coding:utf-8-*-

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

__author__ = 'tony.chenjy'


class DocxGenerator(object):

    def __init__(self):
        pass

    @staticmethod
    def generate_docx(file_path, file_name, tables):
        # make a document
        document = Document()
        document.styles['Normal'].font.name = u'宋体'  # 可换成word里面任意字体

        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(file_name)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.name = u'宋体'

        table_no = 0
        for table in tables:
            print("table: %s columns: %d" % (table, len(table.columns)))

            document.add_page_break()
            table_no += 1

            # document.add_heading(text="%d.%s %s" % (table_no, table.name, table.comment), level=1)
            # document.add_heading(text="%d.%s %s" % (table_no, table.name, table.comment), level=1)
            # document.add_paragraph(text="%s %s" % (table.name, table.comment), style='List Number')
            document.add_heading(text="%s %s" % (table.name, table.comment), level=1)

            trs = len(table.columns) + 1
            # tds = 6
            tds = 5
            t = document.add_table(rows=trs, cols=tds, style='Table Grid')

            # 设置行高
            for i in range(trs):
                tr = t.rows[i]._tr
                trPr = tr.get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), "400")
                trPr.append(trHeight)

            # 设置列宽
            t.columns[0].width = Inches(0.5)
            t.columns[1].width = Inches(1.5)
            t.columns[2].width = Inches(1.1)
            t.columns[3].width = Inches(2)
            t.columns[4].width = Inches(0.9)
            # t.columns[5].width = Inches(1)

            # 设置表格头
            # th_text = [u'序号', u'字段名', u'类型', u'字段描述', u'允许为空', u'默认值']
            th_text = [u'序号', u'字段名', u'类型', u'字段描述', u'允许为空']
            for i in range(len(th_text)):
                p = t.rows[0].cells[i].paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(th_text[i])
                run.font.bold = True
                run.font.size = Pt(12)

            # 设置表格体
            i = 0
            for column in table.columns:
                i += 1
                '''
                for j in range(cols):
                    p = t.rows[i].cells[j].paragraph[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p.add_run(table_headers[i])
                    run.font.bold = True
                    run.font.size = Pt(12)
                '''
                '''
                rowCells = t.rows[i].cells
                rowCells[0].text = str(i)
                rowCells[1].text = column.name
                rowCells[2].text = column.type
                rowCells[3].text = column.comment
                rowCells[4].text = column.allowNull
                if column.defaultValue is not None:
                    rowCells[5].text = column.defaultValue
                '''
                rowCells = t.rows[i].cells
                # 序号
                p = rowCells[0].paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.add_run(str(i))

                rowCells[1].text = column.name
                rowCells[2].text = column.type
                rowCells[3].text = column.comment

                # 是否为空
                p = rowCells[4].paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.add_run(column.allowNull)

                # 默认值
                # if column.defaultValue is not None:
                #     rowCells[5].text = column.defaultValue

        #
        document.save('%s%s.docx' % (file_path, file_name))
        print("Done!")


if __name__ == '__main__':
    m = DocxGenerator()
    print(m)
